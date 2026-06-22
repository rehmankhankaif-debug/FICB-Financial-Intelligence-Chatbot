from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Optional
from uuid import uuid4

from docx import Document

from src.models.document import DocumentChunkSource
from src.utils.errors import IngestionError, UnsupportedFileError
from src.utils.security import get_file_type


def _validate_docx_path(path: Any) -> Path:
    file_path = Path(path)
    if not file_path.exists():
        raise IngestionError("DOCX file does not exist.", metadata={"path": str(file_path)})
    if not file_path.is_file():
        raise IngestionError("DOCX path is not a file.", metadata={"path": str(file_path)})
    if get_file_type(str(file_path)) != "docx":
        raise UnsupportedFileError(
            "DOCX loader only supports .docx files.",
            metadata={"path": str(file_path), "file_type": get_file_type(str(file_path))},
        )
    return file_path


def _extract_tables(document: Document) -> List[str]:
    table_texts: List[str] = []
    for table_index, table in enumerate(document.tables, start=1):
        rows: List[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            table_texts.append("Table {0}:\n{1}".format(table_index, "\n".join(rows)))
    return table_texts


def load_docx(path: Any, source_id: Optional[str] = None) -> List[DocumentChunkSource]:
    file_path = _validate_docx_path(path)
    source_id = source_id or uuid4().hex

    try:
        document = Document(str(file_path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        tables = _extract_tables(document)
        content_parts = paragraphs + tables
        content = "\n\n".join(content_parts).strip()

        metadata: Dict[str, Any] = {
            "source_path": str(file_path),
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
        }

        return [
            DocumentChunkSource(
                source_id=source_id,
                filename=file_path.name,
                source_type="docx",
                content=content,
                metadata=metadata,
            )
        ]
    except Exception as exc:
        raise IngestionError(
            "Failed to load DOCX document.",
            metadata={"path": str(file_path), "error": str(exc)},
        )
