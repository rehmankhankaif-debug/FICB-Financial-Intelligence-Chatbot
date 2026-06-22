from __future__ import annotations

from pathlib import Path

import fitz
import pandas as pd
from docx import Document

from src.agents.query_planner import QueryPlannerAgent
from src.agents.query_rewriter import QueryRewriterAgent
from src.ingestion.docx_loader import load_docx
from src.ingestion.pdf_loader import load_pdf
from src.ingestion.table_loader import load_table
from src.ingestion.url_loader import load_url
from src.llm.gemini_client import GeminiClient
from src.models.query import QueryPlan
from src.table_intelligence.profiler import TableProfiler
from src.tools.chart_tool import ChartTool
from src.tools.compare_tool import CompareTool
from src.tools.rag_qa_tool import RagQATool
from src.tools.summarize_tool import SummarizeTool
from src.tools.table_analysis_tool import TableAnalysisTool


def _dump(model):
    return model.model_dump() if hasattr(model, "model_dump") else model.dict()


def _offline_plan(query: str, source: dict, dataframe: pd.DataFrame | None = None, language: str | None = None):
    client = GeminiClient(api_key="", client=None)
    rewritten = QueryRewriterAgent(client).rewrite(query, language=language)
    profiles = []
    if dataframe is not None:
        profiles = [
            TableProfiler().profile(
                dataframe,
                source_id=source["source_id"],
                filename=source["filename"],
            )
        ]
    return QueryPlannerAgent(client).plan(query, rewritten, [source], profiles), profiles


def _run_table(plan: QueryPlan, dataframe: pd.DataFrame, profile) -> object:
    return TableAnalysisTool().safe_run(
        {"query_plan": _dump(plan), "dataframe": dataframe, "table_profile": _dump(profile)}
    )


def _write_pdf(path: Path, text: str) -> None:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    document.save(path)
    document.close()


def _retrieved_pdf_chunks(path: Path, source_id: str):
    chunks = []
    for index, source in enumerate(load_pdf(path, source_id=source_id), start=1):
        chunks.append(
            {
                "chunk_id": f"{source_id}-{index}",
                "source_id": source_id,
                "filename": path.name,
                "source_type": "pdf",
                "content": source.content,
                "page": index,
                "score": 0.95,
                "metadata": source.metadata,
            }
        )
    return chunks


def test_excel_revenue_trend_executes_sum_and_line_chart(tmp_path: Path) -> None:
    path = tmp_path / "quarterly_revenue.xlsx"
    pd.DataFrame({"Quarter": ["Q1", "Q2"], "Revenue": [100.0, 130.0]}).to_excel(path, index=False)
    dataframe = load_table(path)
    source = {"source_id": "sales", "filename": path.name, "file_type": "xlsx"}
    plan, profiles = _offline_plan(
        "Show me the revenue trends for Q1 and Q2 from the attached Excel file.", source, dataframe
    )

    table_result = _run_table(plan, dataframe, profiles[0])
    chart_result = ChartTool().safe_run(
        {
            "query_plan": _dump(plan),
            "dependency_results": {"table_analysis_tool": _dump(table_result)},
        }
    )

    assert table_result.success is True
    assert {row["Quarter"]: row["sum_Revenue"] for row in table_result.table} == {"Q1": 100.0, "Q2": 130.0}
    assert chart_result.success is True
    assert chart_result.data["chart_type"] == "line"


def test_csv_average_monthly_profit_executes_real_file(tmp_path: Path) -> None:
    path = tmp_path / "profit.csv"
    pd.DataFrame({"Month": ["Jan", "Jan", "Feb"], "Profit": [100, 300, 500]}).to_csv(path, index=False)
    dataframe = load_table(path)
    source = {"source_id": "profit", "filename": path.name, "file_type": "csv"}
    plan, profiles = _offline_plan(
        "What is the average monthly profit according to the CSV data provided?", source, dataframe
    )

    result = _run_table(plan, dataframe, profiles[0])

    assert result.success is True
    assert {row["Month"]: row["mean_Profit"] for row in result.table} == {"Jan": 200.0, "Feb": 500.0}


def test_excel_top_five_products_executes_ranking(tmp_path: Path) -> None:
    path = tmp_path / "products.xlsx"
    pd.DataFrame({"Product": list("ABCDEF"), "Sales": [10, 50, 20, 70, 5, 40]}).to_excel(path, index=False)
    dataframe = load_table(path)
    source = {"source_id": "products", "filename": path.name, "file_type": "xlsx"}
    plan, profiles = _offline_plan(
        "Extract the top five products by sales from the embedded table in the Excel file.", source, dataframe
    )

    result = _run_table(plan, dataframe, profiles[0])

    assert result.success is True
    assert len(result.table) == 5
    assert result.table[0]["Product"] == "D"
    assert result.table[0]["sum_Sales"] == 70


def test_two_real_pdfs_execute_numeric_expense_comparison(tmp_path: Path) -> None:
    report_2023 = tmp_path / "financial_report_2023.pdf"
    report_2022 = tmp_path / "financial_report_2022.pdf"
    _write_pdf(report_2023, "Audited statement: Total expenses were USD 120 million in 2023.")
    _write_pdf(report_2022, "Audited statement: Total expenses were USD 100 million in 2022.")
    query = "Compare the total expenses between the 2023 and 2022 financial reports in the PDF documents."
    query_plan = QueryPlan(
        original_query=query,
        rewritten_query=query,
        intent="compare_documents",
        metrics=[{"name": "expenses"}],
    )
    chunks = _retrieved_pdf_chunks(report_2023, "report_2023") + _retrieved_pdf_chunks(report_2022, "report_2022")

    rag_result = RagQATool().safe_run({"query_plan": _dump(query_plan), "retrieved_chunks": chunks})
    comparison = CompareTool().safe_run(
        {
            "query_plan": _dump(query_plan),
            "dependency_results": {"rag_qa_tool": _dump(rag_result)},
        }
    )

    assert rag_result.success is True
    assert comparison.success is True
    assert comparison.data["absolute_difference"] == 20_000_000.0
    assert comparison.data["percentage_change"] == 20.0
    assert len(comparison.citations) == 2


def test_real_docx_executes_grounded_market_outlook_summary(tmp_path: Path) -> None:
    path = tmp_path / "market_outlook.docx"
    document = Document()
    document.add_heading("Market Outlook", level=1)
    document.add_paragraph("Revenue growth is expected to remain positive while liquidity risk requires monitoring.")
    document.add_table(rows=2, cols=2).cell(0, 0).text = "Metric"
    document.save(path)
    loaded = load_docx(path, source_id="outlook")
    chunks = [
        {
            "chunk_id": "outlook-1",
            "source_id": "outlook",
            "filename": path.name,
            "source_type": "docx",
            "content": loaded[0].content,
            "score": 0.95,
        }
    ]
    plan = QueryPlan(original_query="Summarize the key points from the DOCX document on market outlook.", intent="summarize_document")

    result = SummarizeTool().safe_run({"query_plan": _dump(plan), "document_chunks": chunks})

    assert result.success is True
    assert "revenue growth" in result.answer.lower()
    assert result.citations


def test_spanish_document_query_executes_grounded_spanish_answer() -> None:
    query = "¿Cuáles son las tendencias de ingresos del último trimestre según el informe adjunto?"
    source = {"source_id": "informe", "filename": "informe.pdf", "file_type": "pdf"}
    plan, _ = _offline_plan(query, source, language="es")
    result = RagQATool().safe_run(
        {
            "query_plan": _dump(plan),
            "retrieved_chunks": [
                {
                    "chunk_id": "informe-1",
                    "source_id": "informe",
                    "filename": "informe.pdf",
                    "source_type": "pdf",
                    "page": 3,
                    "score": 0.96,
                    "content": "Los ingresos del último trimestre aumentaron un 12% debido al crecimiento de las ventas.",
                }
            ],
        }
    )

    assert plan.language == "es"
    assert result.success is True
    assert "ingresos" in result.answer.lower()
    assert result.citations[0].page == 3


def test_linked_online_report_executes_fetch_clean_and_grounded_lookup(monkeypatch) -> None:
    class Response:
        status_code = 200
        text = "<html><head><title>Market Report</title></head><body><main>Latest market trends include lower inflation and stronger banking liquidity.</main></body></html>"

        def raise_for_status(self):
            return None

    monkeypatch.setattr("src.ingestion.url_loader.requests.get", lambda *args, **kwargs: Response())
    loaded = load_url("https://example.com/market-report", source_id="market")
    query = "What are the latest market trends mentioned in the linked online report?"
    plan = QueryPlan(original_query=query, rewritten_query=query, intent="url_lookup")
    result = RagQATool().safe_run(
        {
            "query_plan": _dump(plan),
            "retrieved_chunks": [
                {
                    "chunk_id": "market-1",
                    "source_id": "market",
                    "filename": loaded[0].filename,
                    "source_type": "url",
                    "score": 0.95,
                    "content": loaded[0].content,
                }
            ],
        }
    )

    assert "lower inflation" in loaded[0].content.lower()
    assert result.success is True
    assert "banking liquidity" in result.answer.lower()
    assert result.citations
