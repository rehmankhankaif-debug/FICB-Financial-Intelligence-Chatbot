from __future__ import annotations

import pytest
from docx import Document

from src.ingestion.docx_loader import load_docx
from src.utils.errors import IngestionError, UnsupportedFileError


def test_load_docx_extracts_paragraphs_and_tables(tmp_path) -> None:
    path = tmp_path / "outlook.docx"
    document = Document()
    document.add_paragraph("Market outlook remains positive.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Revenue"
    table.cell(1, 1).text = "100"
    document.save(path)

    sources = load_docx(path, source_id="docx_source")

    assert len(sources) == 1
    assert sources[0].source_id == "docx_source"
    assert sources[0].source_type == "docx"
    assert "Market outlook" in sources[0].content
    assert "Revenue | 100" in sources[0].content
    assert sources[0].metadata["table_count"] == 1


def test_load_docx_rejects_wrong_extension(tmp_path) -> None:
    path = tmp_path / "outlook.txt"
    path.write_text("hello", encoding="utf-8")

    with pytest.raises(UnsupportedFileError):
        load_docx(path)


def test_load_docx_handles_corrupted_docx(tmp_path) -> None:
    path = tmp_path / "broken.docx"
    path.write_bytes(b"not a docx")

    with pytest.raises(IngestionError):
        load_docx(path)
